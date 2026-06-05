#!/usr/bin/env python3
"""skeleton_builder.py — deterministic skeleton extractor, zero reasoning.

Input: format_template.docx (any size)
Output: template_skeleton.docx (front matter + styles + header/footer only)

This is the "hands" — deterministic, mechanical transformation.
All semantic reasoning stays with the LLM.

Rules:
  - Scan toàn bộ body._element (không dùng doc.paragraphs)
  - Detect front matter bằng heuristic đơn giản (non-heading có text > 20 chars)
  - Remove body paragraphs/tables sau front matter
  - Guarantee: output có w14:paraId trên mọi paragraph
    (nếu thiếu, inject trước khi save)
  - Preserve: header, footer, styles, page layout, section properties
  - Hash-based cache for repeated template usage
"""
from __future__ import annotations

import hashlib
import json
import shutil
from pathlib import Path
from typing import Any

from lxml import etree
from docx import Document
from docx.oxml.ns import qn


W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

# Keywords used to detect heading styles (mechanical, not semantic)
HEADING_KEYWORDS = [
    "heading", "chapter", "chương", "điều", "mục", "phan", "title",
    "section", "subsection",
]


def get_template_hash(template_path: Path) -> str:
    """SHA-256 of template file for cache invalidation."""
    h = hashlib.sha256()
    h.update(template_path.read_bytes())
    return h.hexdigest()[:16]


def ensure_para_ids(doc: Document) -> int:
    """Inject w14:paraId on paragraphs missing it. Returns count injected."""
    count = 0
    body = doc.element.body
    used_ids: set[str] = set()

    # Collect existing IDs first
    for para_elem in body.iter(qn('w:p')):
        pid = para_elem.get(f"{{{W14}}}paraId")
        if pid:
            used_ids.add(pid)

    # Inject missing
    import random
    for para_elem in body.iter(qn('w:p')):
        pid = para_elem.get(f"{{{W14}}}paraId")
        if pid is None:
            new_id = f"{random.randint(0, 0xFFFFFFFF):08X}"
            while new_id in used_ids:
                new_id = f"{random.randint(0, 0xFFFFFFFF):08X}"
            para_elem.set(f"{{{W14}}}paraId", new_id)
            used_ids.add(new_id)
            count += 1

    return count


def detect_body_start(doc: Document) -> int:
    """Scan body._element directly (p and tbl mixed), return element count where body starts.

    CRITICAL: Must count the same way as the removal loop to ensure index alignment.
    Iterate body._element and count both paragraphs AND tables.

    Heuristic: first non-heading paragraph with text > 20 chars.
    Also checks outline_level_xml for robustness against custom style names.

    This is MECHANICAL detection — LLM does not need to re-reason this.
    """
    from docx.oxml.ns import qn as docx_qn

    body = doc.element.body
    element_count = 0  # Count both p and tbl elements, matching removal loop
    para_index = 0    # Track which doc.paragraphs index we're at

    for child in body:
        tag = etree.QName(child.tag).localname
        if tag == "sectPr":
            break  # Stop at section properties

        # Only count p and tbl elements (skip others like bookmarks, comments, etc.)
        if tag not in ("p", "tbl"):
            continue

        # For paragraphs, check if this is where body starts
        if tag == "p":
            # Match this element to a doc.paragraphs entry
            if para_index < len(doc.paragraphs):
                para = doc.paragraphs[para_index]
                style_name = (para.style.name or "").lower()
                is_heading = any(kw in style_name for kw in HEADING_KEYWORDS)

                # Also check outline level XML for robustness
                if not is_heading:
                    style_el = para.style._element if hasattr(para.style, '_element') else None
                    if style_el is not None:
                        outline_el = style_el.find(docx_qn('w:outlineLvl'))
                        if outline_el is not None:
                            is_heading = True

                # Found first body paragraph — return current element count
                if not is_heading:
                    text = para.text.strip()
                    if len(text) > 20:
                        return element_count

            para_index += 1

        element_count += 1

    return element_count  # entire doc is front matter


def build_skeleton(
    template_path: Path,
    output_path: Path,
    cache_dir: Path | None = None,
    force: bool = False,
) -> dict[str, Any]:
    """Build skeleton from template. Returns metadata dict.

    Args:
        template_path: Path to the source template DOCX.
        output_path: Path where the skeleton DOCX will be written.
        cache_dir: Optional directory for hash-based cache.
        force: If True, skip cache and rebuild.

    Returns:
        Metadata dict with template_hash, body_start_index,
        removed_paragraphs, injected_para_ids, cache_hit.
    """
    template_hash = get_template_hash(template_path)
    cache_meta_path: Path | None = None
    cache_skeleton_path: Path | None = None

    if cache_dir is not None:
        cache_meta_path = cache_dir / f"{template_hash}.meta.json"
        cache_skeleton_path = cache_dir / f"{template_hash}.skeleton.docx"

    # Check cache
    if not force and cache_skeleton_path is not None and cache_skeleton_path.exists() and cache_meta_path is not None and cache_meta_path.exists():
        meta = json.loads(cache_meta_path.read_text())
        shutil.copy2(cache_skeleton_path, output_path)
        meta["cache_hit"] = True
        return meta

    # Cache miss — build skeleton
    doc = Document(str(template_path))
    body = doc.element.body

    # Step 1: detect body start index (full scan)
    body_start_idx = detect_body_start(doc)

    # Step 2: remove body content from body._element
    # Count must match detect_body_start() — count both p and tbl elements
    element_count = 0
    elements_to_remove = []

    for child in list(body):
        tag = etree.QName(child.tag).localname
        if tag == "sectPr":
            break  # preserve section properties

        # Only count and potentially remove p and tbl elements
        if tag in ("p", "tbl"):
            if element_count >= body_start_idx:
                elements_to_remove.append(child)
            element_count += 1  # Increment for EVERY p/tbl, not just p

    removed_count = len(elements_to_remove)
    for el in elements_to_remove:
        body.remove(el)

    # Step 3: ensure ALL remaining paragraphs have w14:paraId
    injected = ensure_para_ids(doc)

    # Step 4: save skeleton
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Step 5: save to cache
    if cache_dir is not None and cache_skeleton_path is not None and cache_meta_path is not None:
        cache_dir.mkdir(parents=True, exist_ok=True)
        shutil.copy2(output_path, cache_skeleton_path)

    meta = {
        "template_hash": template_hash,
        "template_path": str(template_path),
        "body_start_index": body_start_idx,
        "removed_paragraphs": removed_count,
        "injected_para_ids": injected,
        "skeleton_path": str(output_path),
        "cache_hit": False,
    }
    if cache_meta_path is not None:
        cache_meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2))

    print(f"[skeleton_builder] Built skeleton: {removed_count} elements removed, "
          f"{injected} paraIds injected, body_start_index={body_start_idx}")
    return meta


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="Build skeleton DOCX from template.")
    parser.add_argument("--template-file", required=True, help="Path to template DOCX")
    parser.add_argument("--output-file", required=True, help="Path for skeleton output")
    parser.add_argument("--cache-dir", default=None, help="Optional cache directory")
    parser.add_argument("--force", action="store_true", help="Force rebuild, skip cache")
    args = parser.parse_args()

    meta = build_skeleton(
        Path(args.template_file),
        Path(args.output_file),
        cache_dir=Path(args.cache_dir) if args.cache_dir else None,
        force=args.force,
    )
    print(json.dumps(meta, ensure_ascii=False, indent=2))
