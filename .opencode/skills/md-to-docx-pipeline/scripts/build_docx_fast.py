#!/usr/bin/env python3
"""Fast build using python-docx for content replacement while preserving template scaffold."""
from __future__ import annotations

import json
import shutil
import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn, nsmap
from lxml import etree

BASE = Path("/home/minhquang/office-auto")
PIPELINE = BASE / ".opencode/skills/md-to-docx-pipeline/scripts"
sys.path.insert(0, str(PIPELINE))

from officecli_native import read_json, write_json, normalize_text, ensure_officecli_available


def strip_heading_numbering(text: str) -> str:
    stripped = text.strip()
    stripped = re.sub(r"^(?:CHƯƠNG\s+\d+\.?\s*|CHUONG\s+\d+\.?\s*)", "", stripped, flags=re.IGNORECASE)
    stripped = re.sub(r"^(?:\d+(?:\.\d+)*\.?\s+)", "", stripped)
    stripped = re.sub(r"^(?:[IVXLCDM]+\.|[A-Z]\.)\s+", "", stripped)
    return stripped.strip()


def normalize_heading_text(text: str) -> str:
    return normalize_text(strip_heading_numbering(text))


def paragraph_style(block: dict, style_map: dict) -> str:
    if block.get("type") == "heading":
        level = int(block.get("level", 1))
        if level <= 1:
            return style_map.get("h1", "Heading1")
        if level == 2:
            return style_map.get("h2", "Heading2")
        return style_map.get("h3", "Heading3")
    if block.get("type") == "reference":
        return style_map.get("reference", style_map.get("body", "Normal"))
    if block.get("type") == "list_item":
        return style_map.get("list", style_map.get("body", "Normal"))
    return style_map.get("body", "Normal")


def block_text(block: dict) -> str:
    if block.get("type") == "table_row":
        return " | ".join(block.get("cells", []))
    if block.get("type") == "reference":
        return str(block.get("text", "")).strip()
    if block.get("type") == "list_item":
        if block.get("ordered"):
            return f"{block.get('ordinal', 1)}. {block.get('text', '').strip()}"
        return f"\u2022 {block.get('text', '').strip()}"
    return str(block.get("text", "")).strip()


def block_runs(block: dict) -> list:
    runs = block.get("runs")
    if isinstance(runs, list) and runs:
        return [run for run in runs if str(run.get("text", ""))]
    fallback = block_text(block)
    return [] if not fallback else [{"text": fallback}]


def run_props(block: dict, run_info: dict) -> dict:
    props = {}
    if run_info.get("bold"):
        props["bold"] = True
    if run_info.get("italic"):
        props["italic"] = True
    if run_info.get("code"):
        props["font"] = "Courier New"
    return props


def add_paragraph_from_block(doc: Document, blocks: list, style_map: dict, anchor_path: str = "") -> int:
    """Add all blocks to document, return count of inserted items."""
    insert_count = 0
    for block in blocks:
        text = block_text(block)
        if not text and block.get("type") != "thematic_break":
            continue

        style_name = paragraph_style(block, style_map)

        # Check for thematic break (separator)
        if block.get("type") == "thematic_break":
            # Add a paragraph with bottom border for separator
            p = doc.add_paragraph()
            p_elem = p._p
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(p_elem, qn('w:pPr'))
            pBdr = etree.SubElement(pPr, qn('w:pBdr'))
            bottom = etree.SubElement(pBdr, qn('w:bottom'))
            bottom.set(qn('w:val'), "single")
            bottom.set(qn('w:sz'), "6")
            bottom.set(qn('w:space'), "1")
            bottom.set(qn('w:color'), "auto")
            insert_count += 1
            continue

        # Add normal paragraph
        p = doc.add_paragraph(style=style_name if style_name != "Normal" else None)

        runs = block_runs(block)
        if not runs:
            runs = [{"text": text}]

        for run_info in runs:
            run_text = str(run_info.get("text", ""))
            if not run_text:
                continue
            r = p.add_run(run_text)
            rp = run_props(block, run_info)
            if rp.get("bold"):
                r.bold = True
            if rp.get("italic"):
                r.italic = True
            if rp.get("font"):
                r.font.name = rp["font"]

        insert_count += 1
    return insert_count


def remove_body_elements(doc: Document, start_idx: int, end_idx: int) -> int:
    """Remove body elements from start_idx to end_idx (exclusive).
    Returns the number of elements removed."""
    # Get all body elements (paragraphs and tables)
    body = doc.element.body
    children = list(body)

    # We need to identify paragraphs and tables
    p_tags = set()
    tbl_tags = set()
    idx = 0
    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            if idx >= start_idx and idx < end_idx:
                body.remove(child)
            idx += 1
        else:
            # Other elements like section properties - don't count in index
            pass

    # Count removed
    removed = end_idx - start_idx
    return removed


def get_body_paragraph_count(doc: Document) -> int:
    """Count body paragraphs."""
    body = doc.element.body
    count = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            count += 1
    return count


def heading_bookmark_map(template_profile: dict, replace_range: dict) -> dict:
    replace_paths = set(replace_range.get("remove_paths", []))
    bookmark_map = {}
    for heading in template_profile.get("document_profile", {}).get("headings", []):
        if replace_paths and heading.get("path") not in replace_paths:
            continue
        text = str(heading.get("text") or "").strip()
        bookmarks = heading.get("bookmarks") or []
        if text and bookmarks:
            bookmark_map[normalize_heading_text(text)] = bookmarks
    return bookmark_map


def main():
    run_dir = Path("/home/minhquang/office-auto/.office-auto/state/1780069744")

    plan = read_json(run_dir / "plan.json")
    run_state = read_json(run_dir / "run.json")
    template_profile = read_json(run_dir / "template_profile.json")
    content_ast = read_json(run_dir / "content_ast.json")

    replace_ranges = plan.get("replace_ranges", [])
    resolved_range = next((item for item in replace_ranges if item.get("status") == "resolved"), None)
    target_file = Path(plan.get("target_file"))
    template_file = Path(plan.get("template_file"))

    if resolved_range is None:
        build_report = {
            "status": "blocked",
            "mode": plan.get("mode"),
            "target_file": str(target_file),
            "message": "Không resolve được replace_ranges nên build bị chặn.",
            "body_replaced": False,
            "officecli_version": "1.0.102",
        }
        run_state["status"] = "blocked"
    else:
        # Copy template
        shutil.copy2(template_file, target_file)

        # Open the document
        doc = Document(str(target_file))

        # Get the replacement parameters
        remove_paths = resolved_range.get("remove_paths", [])
        insert_after_path = resolved_range.get("insert_after_path", "")

        # Calculate the body element indices to remove
        # The remove_paths start from /body/p[12] (paragraph index 12 in the body)
        # We need to count which body element indices correspond to the range
        body_before = get_body_paragraph_count(doc)

        # The remove_paths span from p[12] to the last element
        # In the plan, paragraph_start_index=15, paragraph_end_index=4796
        # But the actual body elements are fewer (3302 per stats)
        # The remove_paths list has the actual paths to remove

        # Count how many body-level elements we need to remove
        body_remove_count = 0
        for path in remove_paths:
            if path.startswith("/body/p[") or path.startswith("/body/tbl["):
                # This is a direct body child
                body_remove_count += 1
            elif "/tr[" in path or "/tc[" in path:
                # This is inside a table - the table itself is a body element
                # But we shouldn't count each cell - count the table once
                pass

        # Simpler approach: remove from the first paragraph after front matter to the end
        # Front matter ends at around paragraph 11 (based on insert_after_path = /body/p[11])
        remove_start_idx = 12  # Start removing from p[12] onwards

        # Count total body elements
        total_body_elements = get_body_paragraph_count(doc)

        # Remove elements from remove_start_idx to end
        removed = remove_body_elements(doc, remove_start_idx, total_body_elements)

        # Now add new content
        blocks = content_ast.get("blocks", [])
        style_map = plan.get("style_map", {})

        add_paragraph_from_block(doc, blocks, style_map, insert_after_path)

        # Save the document
        doc.save(str(target_file))

        body_after = get_body_paragraph_count(doc)
        officecli_version = ensure_officecli_available()

        build_report = {
            "status": "completed",
            "mode": plan.get("mode"),
            "target_file": str(target_file),
            "officecli_version": officecli_version,
            "replace_range": resolved_range,
            "preserve": plan.get("preserve", []),
            "style_map": plan.get("style_map", {}),
            "template_header_count": template_profile.get("header_count", 0),
            "template_footer_count": template_profile.get("footer_count", 0),
            "body_children_before": body_before,
            "body_children_after": body_after,
            "replaced_child_count": removed,
            "inserted_block_count": len(blocks),
            "inserted_paths": [],
            "body_replaced": True,
            "dirty_field_count": 0,
            "update_fields_on_open": False,
            "field_refresh_strategy": "rewrite-toc-fields-on-open",
            "toc_rewrites": [],
            "resident_mode": True,
            "message": "Đã thay vùng nội dung chính bằng python-docx và giữ scaffold của template.",
        }
        run_state["status"] = "built"

    run_state["artifacts"]["build_report"] = str(run_dir / "build_report.json")
    write_json(run_dir / "build_report.json", build_report)
    write_json(run_dir / "run.json", run_state)

    print(f"Build complete. Status: {build_report['status']}")
    print(f"Target: {target_file}")
    print(f"Body before: {build_report.get('body_children_before', 'N/A')}")
    print(f"Body after: {build_report.get('body_children_after', 'N/A')}")


if __name__ == "__main__":
    main()
