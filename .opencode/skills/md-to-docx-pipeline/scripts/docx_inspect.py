#!/usr/bin/env python3
"""docx_inspect.py — raw dump tool, zero heuristics, zero interpretation.

Schema matches issue.md spec exactly:
  - page_layout_raw: raw twips, no conversion
  - styles_raw: raw pt values (None = inherited), outline_level_xml raw
  - paragraph_sample: first N paragraphs, raw text/style/paraId
  - toc_entries_raw: raw TOC entries from document.xml
  - front_matter_boundary: last paraId before body content zone

LLM receives the raw dump and does ALL reasoning:
  - style inheritance tree resolution
  - spacing interpretation (None = inherited)
  - heading detection from outlineLvl raw values
  - template type classification
  - markdown → style mapping

Usage:
    python docx_inspect.py --template-file <path> --run-dir <path>
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import zipfile
from pathlib import Path
from typing import Any

from lxml import etree


# Word XML namespace
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.openxmlformats.org/markup-compatibility/2006"
A = "http://schemas.openxmlformats.org/drawingml/2006/main"


def timestamp_utc() -> str:
    return dt.datetime.now(dt.timezone.utc).replace(microsecond=0).isoformat()


def get_para_id(para: Any) -> str | None:
    from docx.oxml.ns import qn

    para_id_namespaces = [
        qn('w:paraId'),
        qn('w14:paraId'),
        '{http://schemas.microsoft.com/office/word/2010/wordml}paraId',
    ]

    element = para._element
    # Try as attribute first (for compatibility)
    for ns_qn in para_id_namespaces:
        para_id = element.get(ns_qn)
        if para_id:
            return para_id
    # Try as child element inside w:pPr
    ppr = element.find(qn('w:pPr'))
    if ppr is not None:
        for ns_qn in para_id_namespaces:
            pid_el = ppr.find(ns_qn)
            if pid_el is not None:
                val = pid_el.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                if val:
                    return val
    return None


def collect_theme_fonts(template_path: Path) -> dict[str, str | None]:
    theme_fonts: dict[str, str | None] = {
        "major_latin": None,
        "minor_latin": None,
        "major_ea": None,
        "minor_ea": None,
        "major_cs": None,
        "minor_cs": None,
    }
    try:
        with zipfile.ZipFile(template_path) as zf:
            theme_name = next((name for name in zf.namelist() if name.startswith("word/theme/") and name.endswith(".xml")), None)
            if not theme_name:
                return theme_fonts
            theme_doc = etree.fromstring(zf.read(theme_name))
    except Exception:
        return theme_fonts

    major_font = theme_doc.find(f".//{{{A}}}fontScheme/{{{A}}}majorFont")
    minor_font = theme_doc.find(f".//{{{A}}}fontScheme/{{{A}}}minorFont")

    def read_font(parent: Any, tag: str) -> str | None:
        if parent is None:
            return None
        child = parent.find(f"{{{A}}}{tag}")
        if child is None:
            return None
        return child.get("typeface")

    theme_fonts["major_latin"] = read_font(major_font, "latin")
    theme_fonts["minor_latin"] = read_font(minor_font, "latin")
    theme_fonts["major_ea"] = read_font(major_font, "ea")
    theme_fonts["minor_ea"] = read_font(minor_font, "ea")
    theme_fonts["major_cs"] = read_font(major_font, "cs")
    theme_fonts["minor_cs"] = read_font(minor_font, "cs")
    return theme_fonts


def style_is_heading(style_name: str | None) -> bool:
    if not style_name:
        return False
    lowered = style_name.lower()
    return any(token in lowered for token in ["heading", "chapter", "chương", "điều", "mục", "section", "subsection", "title"])


def resolve_effective_font(para: Any, theme_fonts: dict[str, str | None]) -> str:
    run_candidates: list[str] = []
    for run in getattr(para, "runs", []):
        if getattr(run.font, "name", None):
            run_candidates.append(str(run.font.name))
        if getattr(run.style, "font", None) is not None and getattr(run.style.font, "name", None):
            run_candidates.append(str(run.style.font.name))
        if run_candidates:
            return run_candidates[0]

    style = getattr(para, "style", None)
    visited: set[int] = set()
    while style is not None and id(style) not in visited:
        visited.add(id(style))
        if getattr(style.font, "name", None):
            return str(style.font.name)
        style = getattr(style, "base_style", None)

    style_name = getattr(getattr(para, "style", None), "name", None)
    if style_is_heading(style_name):
        return theme_fonts.get("major_latin") or theme_fonts.get("major_ea") or theme_fonts.get("minor_latin") or "Times New Roman"

    return theme_fonts.get("minor_latin") or theme_fonts.get("minor_ea") or theme_fonts.get("major_latin") or "Times New Roman"


def dump_styles(doc: Any, theme_fonts: dict[str, str | None]) -> list[dict]:
    """Layer 1: Style catalog — RAW, no resolve, no classify.

    For each style, return:
      - name, style_id, base_style name (or None)
      - Raw Pt values (None = inherited, LLM self-computes)
      - outline_level_xml: raw w:outlineLvl @w:val (0-9 = heading levels 1-10)
      - No "heading_level", no "resolved_spacing", no "is_heading"
    """
    from docx.oxml.ns import qn

    styles_raw: list[dict] = []
    for style in doc.styles:
        # Only process PARAGRAPH styles — CHARACTER/TABLE/NUMBERING don't have paragraph_format
        if style.type.name != "PARAGRAPH":
            continue
        fmt = style.paragraph_format
        entry: dict[str, Any] = {
            "name": style.name,
            "style_id": style.style_id,
            "base_style": style.base_style.name if style.base_style else None,
        }
        # Raw Pt values — None means inherited from base style
        if fmt.space_before is not None:
            entry["space_before_pt"] = float(fmt.space_before.pt)
        else:
            entry["space_before_pt"] = None
        if fmt.space_after is not None:
            entry["space_after_pt"] = float(fmt.space_after.pt)
        else:
            entry["space_after_pt"] = None

        if fmt.line_spacing is not None:
            entry["line_spacing"] = float(fmt.line_spacing)
        else:
            entry["line_spacing"] = None

        line_rule = fmt.line_spacing_rule
        if line_rule is not None:
            entry["line_spacing_rule"] = str(line_rule)
        else:
            entry["line_spacing_rule"] = None

        if fmt.left_indent is not None:
            entry["left_indent_pt"] = float(fmt.left_indent.pt)
        else:
            entry["left_indent_pt"] = None

        if fmt.first_line_indent is not None:
            entry["first_line_indent_pt"] = float(fmt.first_line_indent.pt)
        else:
            entry["first_line_indent_pt"] = None

        # Run-level raw style info (no resolve)
        run_info = {}
        if style.font.name is not None:
            run_info["font_name"] = style.font.name
        else:
            run_info["font_name"] = None
        if style.font.size is not None:
            run_info["font_size_pt"] = float(style.font.size.pt)
        else:
            run_info["font_size_pt"] = None
        run_info["font_bold"] = style.font.bold
        run_info["font_italic"] = style.font.italic
        if style.font.color is not None and style.font.color.rgb is not None:
            run_info["font_color_hex"] = str(style.font.color.rgb)
        else:
            run_info["font_color_hex"] = None
        entry["run"] = run_info

        # outline_level_xml: raw w:outlineLvl @w:val from styles.xml
        # 0-9 = heading levels 1-10, None = not a heading
        outline_elem = style._element.find(qn('w:outlineLvl'))
        if outline_elem is not None:
            entry["outline_level_xml"] = outline_elem.get(qn('w:val'))
        else:
            entry["outline_level_xml"] = None

        entry["effective_font"] = resolve_effective_font(style, theme_fonts)

        styles_raw.append(entry)

    return styles_raw


def dump_page_layout(template_path: Path) -> dict[str, str | None]:
    """Layer 2: Page layout — RAW XML values (twips).

    Returns raw twip values from <w:sectPr>/<w:pgSz>/<w:pgMar>.
    LLM converts: twips / 1440 = inches, / 72 = points.
    """
    try:
        with zipfile.ZipFile(template_path) as z:
            xml_doc = etree.fromstring(z.read("word/document.xml"))
    except Exception as exc:
        return {"error": f"Cannot read document.xml: {exc}"}

    sect = xml_doc.find(f".//{{{W}}}sectPr")
    if sect is None:
        return {"error": "No <w:sectPr> found in document.xml"}

    pg_sz = sect.find(f"{{{W}}}pgSz")
    pg_mar = sect.find(f"{{{W}}}pgMar")

    result: dict[str, str | None] = {}

    if pg_sz is not None:
        result["paper_w_twip"] = pg_sz.get(f"{{{W}}}w")
        result["paper_h_twip"] = pg_sz.get(f"{{{W}}}h")
        result["paper_orientation"] = pg_sz.get(f"{{{W}}}orient", "portrait")
    else:
        result["paper_w_twip"] = None
        result["paper_h_twip"] = None
        result["paper_orientation"] = "portrait"

    if pg_mar is not None:
        result["margin_top"] = pg_mar.get(f"{{{W}}}top")
        result["margin_bottom"] = pg_mar.get(f"{{{W}}}bottom")
        result["margin_left"] = pg_mar.get(f"{{{W}}}left")
        result["margin_right"] = pg_mar.get(f"{{{W}}}right")
    else:
        result["margin_top"] = None
        result["margin_bottom"] = None
        result["margin_left"] = None
        result["margin_right"] = None

    return result


def dump_paragraph_sample(doc: Any, max_count: int = 30) -> list[dict]:
    """Layer 3: Sample paragraphs — LLM self-identifies structure.

    Returns raw text + style + paraId. No heading classification.
    Default max_count=30 per issue.md spec.
    """
    from docx.oxml.ns import qn

    # Multiple namespaces that may contain paraId
    PARA_ID_NAMESPACES = [
        qn('w:paraId'),
        qn('w14:paraId'),
        '{http://schemas.microsoft.com/office/word/2010/wordml}paraId',
    ]

    sample_paras: list[dict] = []
    for i, para in enumerate(doc.paragraphs[:max_count]):
        element = para._element

        # Try multiple namespaces for paraId
        para_id = None
        for ns_qn in PARA_ID_NAMESPACES:
            para_id = element.get(ns_qn)
            if para_id:
                break

        entry: dict[str, Any] = {
            "index": i,
            "text": para.text,
            "style_name": para.style.name if para.style else None,
            "style_id": para.style.style_id if para.style else None,
        }
        if para_id is not None:
            entry["para_id"] = para_id

        fmt = para.paragraph_format
        if fmt.space_before is not None:
            entry["space_before_pt"] = float(fmt.space_before.pt)
        else:
            entry["space_before_pt"] = None
        if fmt.space_after is not None:
            entry["space_after_pt"] = float(fmt.space_after.pt)
        else:
            entry["space_after_pt"] = None
        if fmt.left_indent is not None:
            entry["left_indent_pt"] = float(fmt.left_indent.pt)
        else:
            entry["left_indent_pt"] = None
        if fmt.first_line_indent is not None:
            entry["first_line_indent_pt"] = float(fmt.first_line_indent.pt)
        else:
            entry["first_line_indent_pt"] = None

        # Alignment raw (no interpretation)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        align = fmt.alignment
        if align is not None:
            entry["align"] = str(align)
        else:
            entry["align"] = None

        entry["is_empty"] = not para.text.strip()
        sample_paras.append(entry)

    return sample_paras


def dump_all_para_ids(doc: Any, *, front_matter_boundary: dict[str, Any] | None = None, theme_fonts: dict[str, str | None] | None = None) -> list[dict]:
    """Layer 3b: ALL paragraph paraIds — full document, no sampling.

    Iterates body._element directly (not doc.paragraphs) to handle tables
    correctly and provide IDX_ synthetic fallback for paragraphs missing paraId.

    Returns paraId + text preview (first 80 chars) for every paragraph.
    Used by validator to check anchors against the FULL template,
    not just the first 30 paragraphs.

    This is the ground truth for anchor validation.
    """
    from docx.oxml.ns import qn

    PARA_ID_NAMESPACES = [
        qn('w:paraId'),
        qn('w14:paraId'),
        '{http://schemas.microsoft.com/office/word/2010/wordml}paraId',
    ]

    all_paras: list[dict] = []
    front_matter_last_para_id = (front_matter_boundary or {}).get("last_para_id")
    front_matter_body_start = (front_matter_boundary or {}).get("body_start_index")
    front_matter_active = True if front_matter_last_para_id is not None else isinstance(front_matter_body_start, int)

    # Iterate body._element directly (handles tables, paragraphs, etc.)
    body = doc.element.body if hasattr(doc, 'element') and hasattr(doc.element, 'body') else None
    para_index = 0

    if body is not None:
        for child in body:
            tag = etree.QName(child.tag).localname
            if tag == "sectPr":
                break  # stop at section properties
            if tag != "p":
                continue  # skip tables, etc.

            # Try real paraId
            para_id = None
            for ns_qn in PARA_ID_NAMESPACES:
                para_id = child.get(ns_qn)
                if para_id:
                    break

            is_synthetic = False
            if para_id is None:
                # Fallback: synthetic IDX_ paraId
                para_id = f"IDX_{para_index:05d}"
                is_synthetic = True

            # Determine is_front_matter
            is_front_matter = front_matter_active
            if front_matter_last_para_id is None and isinstance(front_matter_body_start, int) and para_index >= int(front_matter_body_start):
                is_front_matter = False

            # Get text preview
            text_content = child.text_content() if hasattr(child, 'text_content') else ""
            text_preview = text_content[:80] if text_content else ""

            # Get style name (may not be available from raw XML)
            style_name = None
            if hasattr(doc, 'paragraphs') and para_index < len(doc.paragraphs):
                style_name = doc.paragraphs[para_index].style.name if doc.paragraphs[para_index].style else None

            # Get effective font
            effective_font = None
            if hasattr(doc, 'paragraphs') and para_index < len(doc.paragraphs):
                effective_font = resolve_effective_font(doc.paragraphs[para_index], theme_fonts or {})

            entry: dict[str, Any] = {
                "index": para_index,
                "para_id": para_id,
                "is_synthetic_id": is_synthetic,
                "text_preview": text_preview,
                "style_name": style_name,
                "effective_font": effective_font,
                "is_front_matter": is_front_matter,
            }
            all_paras.append(entry)

            # Track front_matter transition
            if front_matter_last_para_id and para_id == front_matter_last_para_id:
                front_matter_active = False
            elif front_matter_last_para_id is None and isinstance(front_matter_body_start, int) and para_index >= int(front_matter_body_start):
                front_matter_active = False

            para_index += 1
    else:
        # Fallback to doc.paragraphs if body._element not accessible
        for i, para in enumerate(doc.paragraphs):
            para_id = get_para_id(para)

            if para_id is None:
                para_id = f"IDX_{i:05d}"

            is_front_matter = front_matter_active
            if front_matter_last_para_id is None and isinstance(front_matter_body_start, int) and i >= int(front_matter_body_start):
                is_front_matter = False

            entry: dict[str, Any] = {
                "index": i,
                "para_id": para_id,
                "is_synthetic_id": para_id.startswith("IDX_"),
                "text_preview": (para.text[:80] if para.text else ""),
                "style_name": para.style.name if para.style else None,
                "effective_font": resolve_effective_font(para, theme_fonts or {}),
                "is_front_matter": is_front_matter,
            }
            all_paras.append(entry)

            if front_matter_last_para_id and para_id == front_matter_last_para_id:
                front_matter_active = False
            elif front_matter_last_para_id is None and isinstance(front_matter_body_start, int) and i >= int(front_matter_body_start):
                front_matter_active = False

    return all_paras


def summarize_styles_for_llm(styles_raw: list[dict]) -> dict[str, Any]:
    available_styles: list[dict[str, Any]] = []
    heading_map: dict[str, str] = {}
    body_text_style: str | None = None
    do_not_use_styles = ["Normal", "Default Paragraph Font"]

    for style in styles_raw:
        name = style.get("name")
        style_id = style.get("style_id")
        effective_font = style.get("effective_font") or style.get("run", {}).get("font_name") or style.get("run", {}).get("font")
        if not body_text_style and str(name or "").lower() in {"normal", "normal_style", "body text", "bodytext"}:
            body_text_style = str(style_id or name)

        outline_level = style.get("outline_level_xml")
        if outline_level == "0":
            heading_map.setdefault("chapter_title", str(style_id or name))
        elif outline_level == "1":
            heading_map.setdefault("section", str(style_id or name))
        elif outline_level == "2":
            heading_map.setdefault("subsection", str(style_id or name))

        if str(name or "") in {"Normal", "Default Paragraph Font"} and str(name or "") not in do_not_use_styles:
            do_not_use_styles.append(str(name))

        use_for = "body paragraphs"
        if outline_level is not None:
            use_for = f"heading level {int(outline_level) + 1}"
        elif style_is_heading(str(name)):
            use_for = "heading"

        available_styles.append(
            {
                "name": name,
                "style_id": style_id,
                "use_for": use_for,
                "effective_font": effective_font,
                "font_size_pt": style.get("run", {}).get("font_size_pt"),
                "line_spacing": style.get("line_spacing"),
                "first_line_indent_pt": style.get("first_line_indent_pt"),
            }
        )

    if body_text_style is None and styles_raw:
        first_style = styles_raw[0]
        body_text_style = str(first_style.get("style_id") or first_style.get("name"))

    return {
        "body_text_style": body_text_style,
        "heading_map": heading_map,
        "available_styles": available_styles,
        "do_not_use_styles": do_not_use_styles,
    }


def build_content_map(all_para_ids: list[dict], front_matter_boundary: dict[str, Any]) -> dict[str, Any]:
    front_matter_para_ids = [str(entry.get("para_id")) for entry in all_para_ids if entry.get("is_front_matter")]
    body_placeholder_para_ids = [str(entry.get("para_id")) for entry in all_para_ids if not entry.get("is_front_matter")]
    recommended_insert_anchor = front_matter_boundary.get("last_para_id") or (front_matter_para_ids[-1] if front_matter_para_ids else None)

    return {
        "front_matter": {
            "para_ids": front_matter_para_ids,
            "last_para_id": front_matter_boundary.get("last_para_id"),
            "description": "Title page, TOC, figure list — DO NOT REMOVE",
        },
        "body_placeholders": {
            "para_ids": body_placeholder_para_ids,
            "description": "Placeholder content — SHOULD BE REMOVED before inserting new content",
        },
        "recommended_insert_anchor": recommended_insert_anchor,
    }


def dump_toc_entries(template_path: Path) -> list[dict]:
    """Layer 4: Raw TOC entries from document.xml.

    Extracts TOC fields with their anchor bookmarks and levels.
    No interpretation — raw values only.
    """
    try:
        with zipfile.ZipFile(template_path) as z:
            xml_doc = etree.fromstring(z.read("word/document.xml"))
    except Exception as exc:
        return [{"error": f"Cannot read document.xml: {exc}"}]

    toc_entries: list[dict] = []

    # Find all TOC fields (fldChar with fldCharType="begin")
    for fld_char in xml_doc.findall(f".//{{{W}}}fldChar"):
        fld_char_type = fld_char.get(f"{{{W}}}fldCharType")
        if fld_char_type != "begin":
            continue

        # Get the associated instruction
        instr_text = fld_char.getparent()
        if instr_text is None:
            continue
        parent = instr_text.getparent()
        if parent is None:
            continue

        # Find rInstr (instruction run)
        for r in parent.findall(f"{{{W}}}r"):
            for instr in r.findall(f"{{{W}}}instrText"):
                instruction = instr.text or ""
                if "TOC" not in instruction.upper():
                    continue

                # Extract level from instruction
                level_match = instruction.split(" \\l ")
                level = int(level_match[1].strip()) if len(level_match) > 1 else 1

                # Find bookmark anchor for this TOC entry
                bookmark_anchor = None
                # Look for _Toc bookmarks in the same paragraph
                para = fld_char
                while para is not None and etree.QName(para.tag).localname != "p":
                    para = para.getparent()
                if para is not None:
                    for bookmark_start in para.findall(f".//{{{W}}}bookmark"):
                        name = bookmark_start.get(f"{{{W}}}name")
                        if name and name.startswith("_Toc"):
                            bookmark_anchor = name
                            break

                # Extract text content
                text = ""
                for t in para.findall(f".//{{{W}}}t") if para is not None else []:
                    if t.text:
                        text += t.text[:200]

                if text or bookmark_anchor:
                    toc_entries.append({
                        "text": text.strip(),
                        "bookmark_anchor": bookmark_anchor,
                        "level": level,
                    })

    return toc_entries


def detect_front_matter_boundary(doc: Any) -> dict[str, Any]:
    """Layer 5: Detect front matter boundary by scanning FULL document.

    Heuristic: front matter = content before first body paragraph.
    Body paragraphs typically have "Body Text" or "Normal" style and
    are not headings (outline_level_xml is None or paragraph has no heading style).

    Returns last paraId before body content zone.
    Scans ALL paragraphs in the document, not just a sample.
    """
    from docx.oxml.ns import qn as docx_qn

    # Find first paragraph that looks like body text
    # (not a heading, has substantial text, common body style)
    body_start_index = len(doc.paragraphs)
    for i, para in enumerate(doc.paragraphs):
        style_name = (para.style.name or "").lower()
        text = para.text.strip()
        is_heading = any(kw in style_name for kw in ["heading", "chương", "điều", "mục", "phan"])

        # Also check outline level XML for robustness against custom style names
        if not is_heading:
            style_el = para.style._element if hasattr(para.style, '_element') else None
            if style_el is not None:
                outline_el = style_el.find(docx_qn('w:outlineLvl'))
                if outline_el is not None:
                    is_heading = True

        if not is_heading and len(text) > 20:
            body_start_index = i
            break

    # Last paraId before body
    if body_start_index > 0:
        last_before_body = doc.paragraphs[body_start_index - 1]
        last_para_id = get_para_id(last_before_body)
        return {
            "last_para_id": last_para_id,
            "body_start_index": body_start_index,
            "description": f"body starts at para index {body_start_index} (full scan, {len(doc.paragraphs)} total)",
        }

    return {
        "last_para_id": None,
        "body_start_index": body_start_index,
        "description": "no boundary detected — entire doc may be front matter",
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="docx_inspect: raw dump tool — zero heuristics, zero interpretation."
    )
    parser.add_argument("--template-file", required=True, help="Path to template .docx")
    parser.add_argument("--run-dir", required=True, help="Run directory for output artifacts")
    parser.add_argument(
        "--top-n-styles", type=int, default=None,
        help="Only dump top N styles (by appearance count in sample). "
             "Useful for large templates. This is a FILTERING parameter, not classification."
    )
    parser.add_argument(
        "--max-paragraphs", type=int, default=30,
        help="Number of paragraphs to sample (default: 30 per issue.md spec)"
    )
    parser.add_argument(
        "--skeleton-cache-dir", default=None,
        help="Cache dir for skeleton. If provided, auto-build skeleton before inspect."
    )
    parser.add_argument(
        "--force-skeleton", action="store_true",
        help="Force skeleton rebuild even if cache exists."
    )
    args = parser.parse_args()

    template_path = Path(args.template_file)
    run_dir = Path(args.run_dir)
    run_dir.mkdir(parents=True, exist_ok=True)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Import python-docx here (after validation)
    from docx import Document

    # Skeleton pipeline: auto-build skeleton before inspect if cache dir provided
    if args.skeleton_cache_dir:
        from skeleton_builder import build_skeleton
        skeleton_path = run_dir / "template_skeleton.docx"
        skeleton_meta = build_skeleton(
            template_path,
            skeleton_path,
            cache_dir=Path(args.skeleton_cache_dir),
            force=args.force_skeleton,
        )
        # Use skeleton as input for inspection
        template_path = skeleton_path
        print(f"[docx_inspect] Using skeleton (cache_hit={skeleton_meta['cache_hit']})")

    doc = Document(str(template_path))

    print(f"[docx_inspect] Reading: {template_path}")

    theme_fonts = collect_theme_fonts(template_path)

    # Layer 1: Styles — RAW
    print("[docx_inspect] Layer 1: styles_raw ...")
    styles_raw = dump_styles(doc, theme_fonts)
    if args.top_n_styles and args.top_n_styles > 0:
        # Count style usage in paragraphs
        style_counts: dict[str, int] = {}
        for para in doc.paragraphs:
            sname = para.style.name if para.style else "None"
            style_counts[sname] = style_counts.get(sname, 0) + 1
        most_common = sorted(style_counts.items(), key=lambda x: -x[1])[:args.top_n_styles]
        keep_names = {name for name, _ in most_common}
        styles_raw = [s for s in styles_raw if s["name"] in keep_names]
    Path(run_dir, "docx_inspect_styles_raw.json").write_text(
        json.dumps(styles_raw, ensure_ascii=False, indent=2)
    )
    print(f"[docx_inspect]   → {len(styles_raw)} styles written")

    # Layer 2: Page layout — RAW twips
    print("[docx_inspect] Layer 2: page_layout_raw ...")
    page_layout = dump_page_layout(template_path)
    Path(run_dir, "docx_inspect_page_layout_raw.json").write_text(
        json.dumps(page_layout, ensure_ascii=False, indent=2)
    )
    print("[docx_inspect]   → page_layout written")

    # Layer 3: Sample paragraphs — RAW (default 30 per spec)
    print(f"[docx_inspect] Layer 3: paragraph_sample (max={args.max_paragraphs}) ...")
    para_sample = dump_paragraph_sample(doc, max_count=args.max_paragraphs)
    Path(run_dir, "docx_inspect_paragraph_sample.json").write_text(
        json.dumps(para_sample, ensure_ascii=False, indent=2)
    )
    print(f"[docx_inspect]   → {len(para_sample)} paragraphs written")

    # Layer 4: TOC entries — RAW
    print("[docx_inspect] Layer 4: toc_entries_raw ...")
    toc_entries = dump_toc_entries(template_path)
    Path(run_dir, "docx_inspect_toc_entries_raw.json").write_text(
        json.dumps(toc_entries, ensure_ascii=False, indent=2)
    )
    print(f"[docx_inspect]   → {len(toc_entries)} TOC entries written")

    # Layer 3b: ALL paragraph paraIds — full document (for validator)
    print("[docx_inspect] Layer 3b: all_para_ids (full document) ...")
    front_matter = detect_front_matter_boundary(doc)
    all_para_ids = dump_all_para_ids(doc, front_matter_boundary=front_matter, theme_fonts=theme_fonts)
    Path(run_dir, "docx_inspect_all_para_ids.json").write_text(
        json.dumps(all_para_ids, ensure_ascii=False, indent=2)
    )
    print(f"[docx_inspect]   → {len(all_para_ids)} total paragraphs written")

    styles_for_llm = summarize_styles_for_llm(styles_raw)
    Path(run_dir, "docx_inspect_styles_for_llm.json").write_text(
        json.dumps(styles_for_llm, ensure_ascii=False, indent=2)
    )
    print("[docx_inspect]   → styles_for_llm written")

    # Layer 5: Front matter boundary
    print("[docx_inspect] Layer 5: front_matter_boundary ...")
    Path(run_dir, "docx_inspect_front_matter_boundary.json").write_text(
        json.dumps(front_matter, ensure_ascii=False, indent=2)
    )
    print(f"[docx_inspect]   → front_matter_boundary: {front_matter.get('description')}")

    content_map = build_content_map(all_para_ids, front_matter)
    Path(run_dir, "docx_inspect_content_map.json").write_text(
        json.dumps(content_map, ensure_ascii=False, indent=2)
    )
    print("[docx_inspect]   → content_map written")

    # Build combined output matching issue.md schema exactly
    combined = {
        "source_file": str(template_path.resolve()),
        "extraction_timestamp": timestamp_utc(),
        "page_layout_raw": {
            "paper_w_twip": page_layout.get("paper_w_twip"),
            "paper_h_twip": page_layout.get("paper_h_twip"),
            "margin_top": page_layout.get("margin_top"),
            "margin_bottom": page_layout.get("margin_bottom"),
            "margin_left": page_layout.get("margin_left"),
            "margin_right": page_layout.get("margin_right"),
            "orientation": page_layout.get("paper_orientation", "portrait"),
        },
        "styles_raw": styles_raw,
        "paragraph_sample": para_sample,
        "all_para_ids": all_para_ids,
        "toc_entries_raw": toc_entries,
        "front_matter_boundary": front_matter,
        "styles_for_llm": styles_for_llm,
        "content_map": content_map,
    }
    Path(run_dir, "docx_inspect_output.json").write_text(
        json.dumps(combined, ensure_ascii=False, indent=2)
    )

    # Also write individual layer files for debugging
    summary = {
        "template_file": str(template_path.resolve()),
        "extraction_timestamp": timestamp_utc(),
        "top_n_styles_filter": args.top_n_styles,
        "max_paragraphs": args.max_paragraphs,
        "layer_files": {
            "styles_raw": str(run_dir / "docx_inspect_styles_raw.json"),
            "styles_for_llm": str(run_dir / "docx_inspect_styles_for_llm.json"),
            "page_layout_raw": str(run_dir / "docx_inspect_page_layout_raw.json"),
            "paragraph_sample": str(run_dir / "docx_inspect_paragraph_sample.json"),
            "toc_entries_raw": str(run_dir / "docx_inspect_toc_entries_raw.json"),
            "front_matter_boundary": str(run_dir / "docx_inspect_front_matter_boundary.json"),
            "content_map": str(run_dir / "docx_inspect_content_map.json"),
        },
        "combined_output": str(run_dir / "docx_inspect_output.json"),
    }
    Path(run_dir, "docx_inspect_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2)
    )

    print(f"[docx_inspect] Done. All output in: {run_dir}")
    print("[docx_inspect] NEXT STEP: LLM reads docx_inspect_output.json → reasons → writes execution_ops.json")


if __name__ == "__main__":
    main()
