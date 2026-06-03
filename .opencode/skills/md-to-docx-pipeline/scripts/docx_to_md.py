#!/usr/bin/env python3
"""Extract content from a source DOCX file into Markdown format.

This is a raw extraction tool — zero heuristics, zero interpretation.
It extracts paragraphs, tables, and headings from a DOCX file and
converts them to Markdown for use with the md-to-docx pipeline.

Usage:
    python docx_to_md.py --source-file nghi_dinh.docx --output-md content.md --run-dir .office-auto/state/extract-001/
"""

import argparse
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree


# Namespace mappings for lxml
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
}


def get_para_text(paragraph):
    """Extract plain text from a paragraph."""
    return paragraph.text


def get_para_style(paragraph):
    """Get the style name of a paragraph."""
    if paragraph.style:
        return paragraph.style.name
    return "Normal"


def get_heading_level(paragraph):
    """Extract heading level from style name or outline level."""
    style_name = paragraph.style.name if paragraph.style else ""

    # Check style name for heading patterns
    if "heading" in style_name.lower():
        try:
            # Extract number from "Heading 1", "Heading 2", etc.
            parts = style_name.split()
            for part in parts:
                if part.isdigit():
                    return int(part)
        except (ValueError, IndexError):
            pass

    # Check outline level from XML
    try:
        pPr = paragraph._element.find(".//w:pPr")
        if pPr is not None:
            outlineLvl = pPr.find(".//w:outlineLvl")
            if outlineLvl is not None:
                val = outlineLvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                if val is not None:
                    return int(val) + 1  # 0-based to 1-based
    except Exception:
        pass

    return None


def get_paragraph_alignment(paragraph):
    """Get paragraph alignment."""
    try:
        pPr = paragraph._element.find(".//w:pPr")
        if pPr is not None:
            jc = pPr.find(".//w:jc")
            if jc is not None:
                val = jc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                alignment_map = {
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "both": WD_ALIGN_PARAGRAPH.BOTH,
                    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
                }
                return alignment_map.get(val, WD_ALIGN_PARAGRAPH.LEFT)
    except Exception:
        pass
    return WD_ALIGN_PARAGRAPH.LEFT


def get_run_properties(run):
    """Extract run properties as a dict."""
    props = {}
    if run.bold:
        props["bold"] = True
    if run.italic:
        props["italic"] = True
    if run.underline:
        props["underline"] = True
    if run.font.name:
        props["font"] = run.font.name
    if run.font.size:
        props["size_pt"] = run.font.size.pt
    if run.font.color and run.font.color.rgb:
        props["color"] = run.font.color.rgb
    return props


def para_to_markdown(paragraph):
    """Convert a single paragraph to Markdown."""
    text = paragraph.text.strip()
    if not text:
        return ""

    style = get_para_style(paragraph)
    heading_level = get_heading_level(paragraph)
    alignment = get_paragraph_alignment(paragraph)

    # Build heading prefix
    if heading_level and 1 <= heading_level <= 6:
        prefix = "#" * heading_level + " "
        # Check if text already starts with heading marker
        if text.startswith("#"):
            return text
        return prefix + text

    # Check for list items (bullet/numbered)
    if "bullet" in style.lower() or "list" in style.lower():
        # Check for numbered list
        if "number" in style.lower() or "1." in style.lower() or "1," in style.lower():
            return f"1. {text}"
        return f"- {text}"

    # Check for blockquote
    if "quote" in style.lower() or "citation" in style.lower():
        lines = text.split("\n")
        return "\n".join(f"> {line}" for line in lines)

    # Check alignment
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return f"\n{text}\n"

    return text


def table_to_markdown(table):
    """Convert a table to Markdown format."""
    if not table.rows:
        return ""

    md_lines = []

    # Header row
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    md_lines.append("| " + " | ".join(header_cells) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

    # Data rows
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        md_lines.append("| " + " | ".join(cells) + " |")

    return "\n".join(md_lines)


def extract_content(source_file: Path, max_paragraphs: int = 1000, max_tables: int = 100):
    """Extract content from a DOCX file.

    Returns a dict with:
    - markdown: the full markdown content
    - stats: extraction statistics
    - raw_paragraphs: raw paragraph data for LLM reasoning
    """
    doc = Document(str(source_file))

    md_parts = []
    raw_paragraphs = []
    tables_extracted = 0
    paragraphs_extracted = 0

    for element in doc.element.body:
        tag = etree.QName(element).localname

        if tag == "p" and paragraphs_extracted < max_paragraphs:
            para = doc.paragraphs[paragraphs_extracted]
            text = para.text.strip()

            # Record raw data
            raw_entry = {
                "index": paragraphs_extracted,
                "text": text,
                "style": get_para_style(para),
                "heading_level": get_heading_level(para),
                "alignment": get_paragraph_alignment(para).name if get_paragraph_alignment(para) else "LEFT",
                "runs": [get_run_properties(run) for run in para.runs],
            }
            raw_paragraphs.append(raw_entry)

            # Convert to markdown
            md = para_to_markdown(para)
            if md:
                md_parts.append(md)
            paragraphs_extracted += 1

        elif tag == "tbl" and tables_extracted < max_tables:
            # Find the table index
            table_idx = sum(1 for e in doc.element.body[:doc.element.body.index(element)] if etree.QName(e).localname == "tbl")
            try:
                table = doc.tables[table_idx]
                md = table_to_markdown(table)
                if md:
                    md_parts.append(f"\n{md}\n")
                tables_extracted += 1
            except (IndexError, KeyError):
                pass

    markdown_content = "\n\n".join(md_parts)

    stats = {
        "total_paragraphs": len(doc.paragraphs),
        "extracted_paragraphs": paragraphs_extracted,
        "total_tables": len(doc.tables),
        "extracted_tables": tables_extracted,
        "total_runs": sum(len(p.runs) for p in doc.paragraphs),
    }

    return {
        "markdown": markdown_content,
        "stats": stats,
        "raw_paragraphs": raw_paragraphs,
    }


def main():
    parser = argparse.ArgumentParser(
        description="Extract content from a DOCX file into Markdown format."
    )
    parser.add_argument(
        "--source-file",
        required=True,
        help="Path to the source DOCX file to extract from.",
    )
    parser.add_argument(
        "--output-md",
        default=None,
        help="Path to write the extracted Markdown. Defaults to <run_dir>/extracted_content.md.",
    )
    parser.add_argument(
        "--output-raw",
        default=None,
        help="Path to write the raw paragraph data. Defaults to <run_dir>/extracted_raw.json.",
    )
    parser.add_argument(
        "--run-dir",
        default=".office-auto/state/extract-001",
        help="Run directory for output artifacts.",
    )
    parser.add_argument(
        "--max-paragraphs",
        type=int,
        default=1000,
        help="Maximum number of paragraphs to extract (default: 1000).",
    )
    parser.add_argument(
        "--max-tables",
        type=int,
        default=100,
        help="Maximum number of tables to extract (default: 100).",
    )

    args = parser.parse_args()

    source = Path(args.source_file)
    if not source.exists():
        print(f"Error: Source file not found: {source}", file=sys.stderr)
        sys.exit(1)

    run_dir = Path(args.run_dir)
    run_dir.mkdir(parents=True, exist_ok=True)

    # Determine output paths
    output_md = Path(args.output_md) if args.output_md else run_dir / "extracted_content.md"
    output_raw = Path(args.output_raw) if args.output_raw else run_dir / "extracted_raw.json"

    # Extract content
    print(f"Extracting content from: {source}")
    result = extract_content(source, args.max_paragraphs, args.max_tables)

    # Write markdown
    with open(output_md, "w", encoding="utf-8") as f:
        f.write(result["markdown"])
    print(f"Markdown written to: {output_md}")

    # Write raw data
    raw_output = {
        "source_file": str(source),
        "extraction_timestamp": datetime.now(timezone.utc).isoformat(),
        "stats": result["stats"],
        "raw_paragraphs": result["raw_paragraphs"],
    }
    with open(output_raw, "w", encoding="utf-8") as f:
        json.dump(raw_output, f, indent=2, ensure_ascii=False)
    print(f"Raw data written to: {output_raw}")

    # Print summary
    print(f"\nExtraction Summary:")
    print(f"  Paragraphs: {result['stats']['extracted_paragraphs']}/{result['stats']['total_paragraphs']}")
    print(f"  Tables: {result['stats']['extracted_tables']}/{result['stats']['total_tables']}")
    print(f"  Markdown size: {len(result['markdown'])} characters")


if __name__ == "__main__":
    main()
