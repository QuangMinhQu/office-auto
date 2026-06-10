from __future__ import annotations

import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch


REPO_ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = REPO_ROOT / ".opencode/skills/md-to-docx-pipeline/scripts"
sys.path.insert(0, str(SCRIPTS_DIR))
sys.path.insert(0, str(REPO_ROOT / "scripts"))

import build_report
import docx_validate_ops
import officecli_native
import skeleton_builder


class ValidateOpsTests(unittest.TestCase):
    def test_validate_ops_payload_warns_for_unknown_style_and_anchor(self) -> None:
        template_inspection = {
            "styles_raw": [{"style_id": "Normal"}, {"style_id": "Heading1"}],
            "all_para_ids": [
                {"para_id": "001", "is_front_matter": True},
                {"para_id": "002", "is_front_matter": True},
            ],
            "paragraph_sample": [{"para_id": "001"}],
            "body_children": [{"path": "/body/p[1]"}],
            "body_paragraphs": [{"path": "/body/p[1]"}],
        }
        ops_payload = {
            "ops": [
                {
                    "op": "insert_paragraph_after",
                    "anchor": "/body/p[999]",
                    "style": "MissingStyle",
                    "text": "Test",
                }
            ]
        }

        warnings = docx_validate_ops.validate_ops_payload(ops_payload, template_inspection)

        self.assertEqual(len(warnings), 2)
        self.assertTrue(any(w["type"] == "unknown_anchor" for w in warnings))
        self.assertTrue(any(w["type"] == "unknown_style" for w in warnings))


class BuildReportTests(unittest.TestCase):
    def test_officecli_version_returns_text(self) -> None:
        version = build_report.officecli_version()

        self.assertIsInstance(version, str)
        self.assertTrue(version)

    def test_write_json_roundtrip(self) -> None:
        run_dir = REPO_ROOT / ".office-auto" / "state" / "test-write-json"
        run_dir.mkdir(parents=True, exist_ok=True)
        try:
            target = run_dir / "sample.json"
            build_report.write_json(target, {"a": 1})
            self.assertEqual(officecli_native.read_json(target), {"a": 1})
        finally:
            for path in sorted(run_dir.glob("*"), reverse=True):
                path.unlink(missing_ok=True)
            run_dir.rmdir()


class SkeletonBuilderTests(unittest.TestCase):
    def test_get_template_hash_returns_16_char_hex(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(b"test content")
            f.flush()
            try:
                h = skeleton_builder.get_template_hash(Path(f.name))
                self.assertEqual(len(h), 16)
                int(h, 16)  # should not raise
            finally:
                Path(f.name).unlink()

    def test_build_skeleton_creates_output(self) -> None:
        # Create a proper DOCX using python-docx
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "template.docx"
            output_path = tmpdir_path / "skeleton.docx"
            cache_dir = tmpdir_path / "cache"

            # Create a proper DOCX with front matter content
            doc = Document()
            doc.add_heading("Title Page", level=0)
            doc.add_paragraph("This is the title page content.")
            doc.save(str(template_path))

            meta = skeleton_builder.build_skeleton(template_path, output_path, cache_dir=cache_dir)

            self.assertTrue(output_path.exists())
            self.assertEqual(meta["cache_hit"], False)
            self.assertIn("template_hash", meta)
            self.assertIn("body_start_index", meta)
            self.assertIn("removed_paragraphs", meta)
            self.assertIn("injected_para_ids", meta)

    def test_build_skeleton_cache_hit(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "template.docx"
            output_path = tmpdir_path / "skeleton.docx"
            cache_dir = tmpdir_path / "cache"

            # Create a proper DOCX
            doc = Document()
            doc.add_heading("Title Page", level=0)
            doc.add_paragraph("This is the title page content.")
            doc.save(str(template_path))

            # First build
            meta1 = skeleton_builder.build_skeleton(template_path, output_path, cache_dir=cache_dir)
            self.assertEqual(meta1["cache_hit"], False)

            # Second build (should hit cache)
            output_path.unlink()
            meta2 = skeleton_builder.build_skeleton(template_path, output_path, cache_dir=cache_dir)
            self.assertEqual(meta2["cache_hit"], True)
            self.assertTrue(output_path.exists())

    def test_ensure_para_ids_injects_missing(self) -> None:
        from docx import Document
        from docx.oxml.ns import qn

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            template_path = tmpdir_path / "template.docx"
            output_path = tmpdir_path / "output.docx"

            # Create a DOCX with a paragraph that has no paraId
            doc = Document()
            doc.add_paragraph("Test paragraph")
            doc.save(str(template_path))

            # Reload and remove paraId
            doc = Document(str(template_path))
            body = doc.element.body
            for para_elem in list(body.iter(qn('w:p'))):
                para_elem.attrib.pop('{http://schemas.microsoft.com/office/word/2010/wordml}paraId', None)

            doc.save(str(output_path))

            # Reload and test ensure_para_ids
            doc2 = Document(str(output_path))
            body2 = doc2.element.body
            injected = skeleton_builder.ensure_para_ids(doc2)
            self.assertGreaterEqual(injected, 1)

            # Verify paraId was injected (check on doc2's body, not doc's)
            for para_elem in body2.iter(qn('w:p')):
                pid = para_elem.get('{http://schemas.microsoft.com/office/word/2010/wordml}paraId')
                self.assertIsNotNone(pid)


class ResolveIdxAnchorTests(unittest.TestCase):
    def test_resolve_idx_anchor_returns_none_for_non_idx(self) -> None:
        from execute_execution_ops import resolve_idx_anchor
        result = resolve_idx_anchor("/body/p[@paraId=ABCD]", MagicMock())
        self.assertIsNone(result)

    def test_resolve_idx_anchor_returns_none_for_invalid_format(self) -> None:
        from execute_execution_ops import resolve_idx_anchor
        result = resolve_idx_anchor("IDX_NOTANUMBER", MagicMock())
        self.assertIsNone(result)


class ValidateOpsIdxAwarenessTests(unittest.TestCase):
    def test_has_synthetic_para_ids_returns_true(self) -> None:
        template_inspection = {
            "all_para_ids": [
                {"para_id": "001"},
                {"para_id": "IDX_00005"},
            ],
        }
        self.assertTrue(docx_validate_ops.has_synthetic_para_ids(template_inspection))

    def test_has_synthetic_para_ids_returns_false(self) -> None:
        template_inspection = {
            "all_para_ids": [
                {"para_id": "001"},
                {"para_id": "002"},
            ],
        }
        self.assertFalse(docx_validate_ops.has_synthetic_para_ids(template_inspection))

    def test_resolve_anchor_for_op_accepts_idx_when_synthetic(self) -> None:
        known_paths: set[str] = set()
        known_para_ids: set[str] = {"001", "002"}
        operation = {"anchor": "IDX_00005"}
        resolved, is_valid = docx_validate_ops.resolve_anchor_for_op(
            operation, known_paths, known_para_ids, has_synthetic=True
        )
        self.assertTrue(is_valid)
        self.assertEqual(resolved, "IDX_00005")

    def test_resolve_anchor_for_op_rejects_idx_when_not_synthetic(self) -> None:
        known_paths: set[str] = set()
        known_para_ids: set[str] = {"001", "002"}
        operation = {"anchor": "IDX_00005"}
        resolved, is_valid = docx_validate_ops.resolve_anchor_for_op(
            operation, known_paths, known_para_ids, has_synthetic=False
        )
        self.assertFalse(is_valid)

    def test_validate_ops_with_synthetic_anchor_no_high_warning(self) -> None:
        template_inspection = {
            "styles_raw": [{"style_id": "Normal"}],
            "all_para_ids": [
                {"para_id": "001", "is_front_matter": True},
                {"para_id": "IDX_00005", "is_synthetic_id": True, "is_front_matter": True},
            ],
            "paragraph_sample": [{"para_id": "001"}],
            "body_children": [{"path": "/body/p[1]"}],
            "body_paragraphs": [{"path": "/body/p[1]"}],
        }
        ops_payload = {
            "ops": [
                {
                    "op": "insert_paragraph_after",
                    "anchor": "IDX_00005",
                    "style": "Normal",
                    "text": "Test",
                    "role": "body",
                }
            ]
        }
        warnings = docx_validate_ops.validate_ops_payload(ops_payload, template_inspection)
        # Should have no HIGH warnings (completeness warnings about missing remove ops are expected
        # in this narrow test, but the synthetic anchor and style should be valid)
        high_warnings = [w for w in warnings if w.get("severity") == "high"]
        self.assertEqual(len(high_warnings), 0, f"Expected no high warnings, got: {high_warnings}")


class SafeXmlToStringTests(unittest.TestCase):
    def test_safe_xml_tostring_raises_without_lxml(self) -> None:
        # This test verifies the error message when lxml is not available
        # In practice, lxml IS available, so we test the error path by mocking
        with patch.object(officecli_native, '_USE_LXML', False):
            with self.assertRaises(RuntimeError) as ctx:
                officecli_native.safe_xml_tostring(MagicMock())
            self.assertIn("lxml is required", str(ctx.exception))


class ExecuteExecutionOpsFailFastTests(unittest.TestCase):
    """Test --fail-fast mode stops immediately on first error instead of accumulating errors."""

    def test_fail_fast_false_accumulates_errors(self) -> None:
        """When fail_fast=False, errors are accumulated and execution continues."""
        import execute_execution_ops

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            target_path = tmpdir_path / "test.docx"

            # Create a simple DOCX
            from docx import Document
            doc = Document()
            doc.add_paragraph("Test content")
            doc.save(str(target_path))

            # Mock officecli operations to fail on specific ops
            ops = [
                {"op": "insert_paragraph_after", "anchor": "/body/p[1]", "style": "Normal", "text": "First"},
                {"op": "insert_paragraph_after", "anchor": "/body/p[2]", "style": "Normal", "text": "Second"},
                {"op": "insert_paragraph_after", "anchor": "/body/p[3]", "style": "Normal", "text": "Third"},
            ]

            with patch("execute_execution_ops.officecli_open"):
                with patch("execute_execution_ops.officecli_close"):
                    with patch("execute_execution_ops.officecli_save"):
                        with patch("execute_execution_ops.execute_insert_paragraph", side_effect=officecli_native.OfficeCliError("Mock error")):
                            report = execute_execution_ops.execute_ops_batch(
                                ops,
                                target_path,
                                fail_fast=False
                            )

            # With fail_fast=False, all 3 ops fail but execution completes
            self.assertEqual(report["failed"], 3)
            self.assertEqual(report["total_ops"], 3)
            self.assertEqual(len(report["errors"]), 3)

    def test_fail_fast_true_stops_on_first_error(self) -> None:
        """When fail_fast=True, execution stops immediately on first error."""
        import execute_execution_ops

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            target_path = tmpdir_path / "test.docx"

            # Create a simple DOCX
            from docx import Document
            doc = Document()
            doc.add_paragraph("Test content")
            doc.save(str(target_path))

            # Mock officecli operations to fail on first op
            ops = [
                {"op": "insert_paragraph_after", "anchor": "/body/p[1]", "style": "Normal", "text": "First"},
                {"op": "insert_paragraph_after", "anchor": "/body/p[2]", "style": "Normal", "text": "Second"},
                {"op": "insert_paragraph_after", "anchor": "/body/p[3]", "style": "Normal", "text": "Third"},
            ]

            with patch("execute_execution_ops.officecli_open"):
                with patch("execute_execution_ops.officecli_close"):
                    with patch("execute_execution_ops.officecli_save"):
                        with patch("execute_execution_ops.execute_insert_paragraph", side_effect=officecli_native.OfficeCliError("Mock error")):
                            # With fail_fast=True, should raise on first error
                            with self.assertRaises(officecli_native.OfficeCliError):
                                execute_execution_ops.execute_ops_batch(
                                    ops,
                                    target_path,
                                    fail_fast=True
                                )


if __name__ == "__main__":
    unittest.main()
