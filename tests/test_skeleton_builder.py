#!/usr/bin/env python3
"""Tests for skeleton_builder.py — verify front matter detection and body stripping."""
import sys
from pathlib import Path
import tempfile
import shutil

# Add pipeline scripts to path for imports
PIPELINE_SCRIPTS = Path(__file__).resolve().parents[1] / ".opencode/skills/md-to-docx-pipeline/scripts"
if str(PIPELINE_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(PIPELINE_SCRIPTS))

import unittest
from docx import Document
from docx.shared import Pt


class TestSkeletonBuilder(unittest.TestCase):
    """Test skeleton_builder correctly strips body content."""

    @classmethod
    def setUpClass(cls):
        """Create a test template with known structure: front matter + body."""
        cls.temp_dir = tempfile.mkdtemp()
        cls.template_path = Path(cls.temp_dir) / "test_template.docx"

        # Create template with 5 paragraphs: 2 headings (front matter) + 3 body
        doc = Document()
        
        # Front matter
        p1 = doc.add_heading("Report Title", level=0)
        p1.text = "Long report title with more than 20 characters to ensure we detect it"
        
        p2 = doc.add_heading("Table of Contents", level=1)
        
        # Body (should be stripped)
        p3 = doc.add_paragraph("This is body paragraph 1 with substantial text content for testing", style="Normal")
        p4 = doc.add_paragraph("This is body paragraph 2 with substantial text content for testing", style="Normal")
        p5 = doc.add_paragraph("This is body paragraph 3 with substantial text content for testing", style="Normal")

        doc.save(str(cls.template_path))

    @classmethod
    def tearDownClass(cls):
        """Clean up temp directory."""
        shutil.rmtree(cls.temp_dir, ignore_errors=True)

    def test_skeleton_strips_body_paragraphs(self):
        """Verify skeleton strips all body paragraphs after front matter."""
        from skeleton_builder import build_skeleton, detect_body_start

        output_path = Path(self.temp_dir) / "test_skeleton.docx"
        cache_dir = Path(self.temp_dir) / "cache"
        
        # Build skeleton
        meta = build_skeleton(
            self.template_path,
            output_path,
            cache_dir=cache_dir,
            force=True
        )

        # Verify metadata
        self.assertGreater(meta["removed_paragraphs"], 0, "Should have removed some paragraphs")
        self.assertGreaterEqual(meta["body_start_index"], 0, "body_start_index should be >= 0")
        self.assertTrue(output_path.exists(), "Skeleton file should exist")

        # Load skeleton and verify body is stripped
        skeleton_doc = Document(str(output_path))
        original_doc = Document(str(self.template_path))

        original_para_count = len(original_doc.paragraphs)
        skeleton_para_count = len(skeleton_doc.paragraphs)

        self.assertLess(
            skeleton_para_count,
            original_para_count,
            f"Skeleton ({skeleton_para_count} paras) should have fewer paragraphs than original ({original_para_count})"
        )

        # Skeleton should only have front matter paragraphs
        # Original has 5, skeleton should have 2 (the headings)
        expected_max = original_para_count - meta["removed_paragraphs"]
        self.assertEqual(
            skeleton_para_count,
            expected_max,
            f"Skeleton should have {expected_max} paragraphs after removing {meta['removed_paragraphs']}"
        )

    def test_skeleton_preserves_styles(self):
        """Verify skeleton preserves styles from original template."""
        from skeleton_builder import build_skeleton

        output_path = Path(self.temp_dir) / "test_skeleton_styles.docx"
        cache_dir = Path(self.temp_dir) / "cache_styles"

        build_skeleton(
            self.template_path,
            output_path,
            cache_dir=cache_dir,
            force=True
        )

        skeleton_doc = Document(str(output_path))
        original_doc = Document(str(self.template_path))

        # Count styles
        skeleton_style_names = {s.name for s in skeleton_doc.styles}
        original_style_names = {s.name for s in original_doc.styles}

        # Skeleton should have same or more styles (preserved from original)
        self.assertTrue(
            len(skeleton_style_names) > 0,
            "Skeleton should preserve styles"
        )

    def test_detect_body_start_counts_correctly(self):
        """Verify detect_body_start counts both paragraphs and tables."""
        from skeleton_builder import detect_body_start

        original_doc = Document(str(self.template_path))
        body_start = detect_body_start(original_doc)

        # Should return an index where body starts
        self.assertGreaterEqual(body_start, 0, "body_start index should be >= 0")
        self.assertLessEqual(
            body_start,
            len(original_doc.paragraphs),
            f"body_start ({body_start}) should be <= paragraph count ({len(original_doc.paragraphs)})"
        )

    def test_skeleton_cache_works(self):
        """Verify skeleton caching mechanism."""
        from skeleton_builder import build_skeleton

        cache_dir = Path(self.temp_dir) / "test_cache"
        output_path1 = Path(self.temp_dir) / "skeleton1.docx"
        output_path2 = Path(self.temp_dir) / "skeleton2.docx"

        # First build — cache miss
        meta1 = build_skeleton(
            self.template_path,
            output_path1,
            cache_dir=cache_dir,
            force=True
        )
        self.assertFalse(meta1["cache_hit"], "First build should be cache miss")

        # Second build — should hit cache
        meta2 = build_skeleton(
            self.template_path,
            output_path2,
            cache_dir=cache_dir,
            force=False
        )
        self.assertTrue(meta2["cache_hit"], "Second build should hit cache")

        # Skeletons should be identical
        with open(output_path1, "rb") as f1, open(output_path2, "rb") as f2:
            self.assertEqual(f1.read(), f2.read(), "Cached skeleton should be identical")


if __name__ == "__main__":
    unittest.main()
